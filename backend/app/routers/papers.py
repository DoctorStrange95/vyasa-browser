"""
Paper writing router.

POST   /api/papers                          — create new paper
GET    /api/papers                          — list user's papers
GET    /api/papers/{id}                     — get paper (with content)
PUT    /api/papers/{id}                     — autosave content
DELETE /api/papers/{id}                     — delete paper

POST   /api/papers/{id}/references          — add reference
DELETE /api/papers/{id}/references/{ref_id} — remove reference
GET    /api/papers/{id}/references          — list references

POST   /api/papers/{id}/generate-section    — AI section generator
POST   /api/papers/suggest-references       — find relevant articles by paragraph
POST   /api/papers/{id}/check-abstract      — abstract checker

GET    /api/papers/{id}/export              — export ?format=docx|ris
"""
from __future__ import annotations

import io
import json
import os
from datetime import datetime, timezone
from typing import Annotated, Any

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.models import Article, Paper, PaperStatus, Reference
from app.routers.auth import CurrentUser

router = APIRouter(prefix="/api/papers", tags=["papers"])

# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class PaperCreate(BaseModel):
    title: str
    abstract: str | None = None


class PaperUpdate(BaseModel):
    title: str | None = None
    abstract: str | None = None
    content: dict | None = None   # TipTap JSON document
    status: PaperStatus | None = None


class PaperSummary(BaseModel):
    id: int
    title: str
    abstract: str | None
    status: PaperStatus
    doi: str | None
    created_at: datetime
    updated_at: datetime

    model_config = {"from_attributes": True}


class PaperDetail(PaperSummary):
    content: dict | None = None


class ReferenceCreate(BaseModel):
    article_id: int
    citation_style: str = "vancouver"
    position: int | None = None


class ReferenceResponse(BaseModel):
    id: int
    article_id: int
    citation_style: str
    citation_text: str | None
    position: int | None

    model_config = {"from_attributes": True}


class GenerateSectionRequest(BaseModel):
    section: str   # introduction | methods | results | discussion | conclusion
    context: dict  # {title, abstract, references: [{title, authors, year, journal}]}


class SuggestReferencesRequest(BaseModel):
    paragraph_text: str
    limit: int = 10


class CheckAbstractRequest(BaseModel):
    abstract_text: str
    paper_content: dict | None = None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _format_vancouver(article: Article, position: int) -> str:
    """Minimal Vancouver citation."""
    authors = ", ".join((article.authors or [])[:3])
    if len(article.authors or []) > 3:
        authors += " et al"
    journal = article.journal or ""
    year = article.year or ""
    vol = f";{article.volume}" if article.volume else ""
    pages = f":{article.pages}" if article.pages else ""
    doi = f" doi:{article.doi}" if article.doi else ""
    return f"{position}. {authors}. {article.title}. {journal}. {year}{vol}{pages}.{doi}"


def _format_apa(article: Article) -> str:
    authors_raw = article.authors or []
    if len(authors_raw) == 1:
        author_str = authors_raw[0]
    elif len(authors_raw) <= 6:
        author_str = ", ".join(authors_raw[:-1]) + f", & {authors_raw[-1]}"
    else:
        author_str = ", ".join(authors_raw[:6]) + " et al."
    year = f"({article.year})" if article.year else "(n.d.)"
    journal = article.journal or ""
    vol = article.volume or ""
    pages = article.pages or ""
    doi = f" https://doi.org/{article.doi}" if article.doi else ""
    return f"{author_str} {year}. {article.title}. {journal}, {vol}, {pages}.{doi}"


def _format_harvard(article: Article) -> str:
    authors_raw = article.authors or []
    if len(authors_raw) == 1:
        author_str = authors_raw[0]
    elif len(authors_raw) <= 3:
        author_str = ", ".join(authors_raw[:-1]) + f" and {authors_raw[-1]}"
    else:
        author_str = authors_raw[0] + " et al."
    year = article.year or "n.d."
    journal = article.journal or ""
    vol = f", {article.volume}" if article.volume else ""
    issue = f"({article.issue})" if article.issue else ""
    pages = f", pp. {article.pages}" if article.pages else ""
    doi = f". doi:{article.doi}" if article.doi else ""
    return f"{author_str} ({year}) '{article.title}', {journal}{vol}{issue}{pages}{doi}."


def _format_mla(article: Article) -> str:
    authors_raw = article.authors or []
    if not authors_raw:
        author_str = "Unknown"
    elif len(authors_raw) == 1:
        author_str = authors_raw[0]
    else:
        author_str = f"{authors_raw[0]}, et al."
    year = article.year or "n.d."
    journal = article.journal or ""
    vol = article.volume or ""
    pages = article.pages or ""
    doi = f" doi:{article.doi}" if article.doi else ""
    return f'{author_str}. "{article.title}." {journal}, vol. {vol}, {year}, pp. {pages}.{doi}'


def _format_chicago(article: Article) -> str:
    authors_raw = article.authors or []
    if not authors_raw:
        author_str = "Unknown"
    elif len(authors_raw) <= 3:
        author_str = ", ".join(authors_raw)
    else:
        author_str = f"{authors_raw[0]} et al."
    year = article.year or "n.d."
    journal = article.journal or ""
    vol = article.volume or ""
    issue = f", no. {article.issue}" if article.issue else ""
    pages = f": {article.pages}" if article.pages else ""
    doi = f". https://doi.org/{article.doi}" if article.doi else ""
    return f'{author_str}. "{article.title}." {journal} {vol}{issue} ({year}){pages}{doi}.'


def _format_ama(article: Article, position: int) -> str:
    """American Medical Association (AMA) 11th edition."""
    authors_raw = article.authors or []
    if len(authors_raw) > 6:
        author_str = ", ".join(authors_raw[:6]) + ", et al."
    else:
        author_str = ", ".join(authors_raw)
    year = article.year or "n.d."
    journal = article.journal or ""
    vol = article.volume or ""
    issue = f"({article.issue})" if article.issue else ""
    pages = f":{article.pages}" if article.pages else ""
    doi = f". doi:{article.doi}" if article.doi else ""
    return f"{position}. {author_str}. {article.title}. {journal}. {year};{vol}{issue}{pages}{doi}"


def _format_nature(article: Article, position: int) -> str:
    """Nature journal citation style."""
    authors_raw = article.authors or []
    if len(authors_raw) > 5:
        author_str = ", ".join(authors_raw[:5]) + " et al."
    else:
        author_str = ", ".join(authors_raw)
    year = article.year or "n.d."
    journal = article.journal or ""
    vol = article.volume or ""
    pages = article.pages or ""
    doi = f" https://doi.org/{article.doi}" if article.doi else ""
    return f"{position}. {author_str}. {article.title}. {journal} {vol}, {pages} ({year}){doi}."


def _format_citation(article: Article, style: str, position: int) -> str:
    s = style.lower()
    if s == "apa":
        return _format_apa(article)
    if s == "harvard":
        return _format_harvard(article)
    if s == "mla":
        return _format_mla(article)
    if s == "chicago":
        return _format_chicago(article)
    if s == "ama":
        return _format_ama(article, position)
    if s == "nature":
        return _format_nature(article, position)
    # Default: Vancouver
    return _format_vancouver(article, position)


# ---------------------------------------------------------------------------
# PAPER CRUD
# ---------------------------------------------------------------------------

@router.post("", response_model=PaperSummary, status_code=status.HTTP_201_CREATED)
async def create_paper(
    body: PaperCreate,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> Paper:
    paper = Paper(
        user_id=current_user.id,
        title=body.title,
        abstract=body.abstract,
        content={
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": body.title}]},
                {"type": "paragraph", "content": [{"type": "text", "text": ""}]},
            ],
        },
        status=PaperStatus.DRAFT,
    )
    db.add(paper)
    await db.flush()
    return paper


@router.get("", response_model=list[PaperSummary])
async def list_papers(
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> list[Paper]:
    result = await db.execute(
        select(Paper)
        .where(Paper.user_id == current_user.id)
        .order_by(Paper.updated_at.desc())
    )
    return list(result.scalars().all())


@router.get("/{paper_id}", response_model=PaperDetail)
async def get_paper(
    paper_id: int,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> Paper:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    return paper


@router.put("/{paper_id}", response_model=PaperSummary)
async def update_paper(
    paper_id: int,
    body: PaperUpdate,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> Paper:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    if body.title is not None:
        paper.title = body.title
    if body.abstract is not None:
        paper.abstract = body.abstract
    if body.content is not None:
        paper.content = body.content
    if body.status is not None:
        paper.status = body.status
    paper.updated_at = datetime.now(timezone.utc)

    return paper


@router.delete("/{paper_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_paper(
    paper_id: int,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> None:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    await db.delete(paper)


# ---------------------------------------------------------------------------
# REFERENCES
# ---------------------------------------------------------------------------

@router.post("/{paper_id}/references", response_model=ReferenceResponse, status_code=status.HTTP_201_CREATED)
async def add_reference(
    paper_id: int,
    body: ReferenceCreate,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> Reference:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    article = await db.get(Article, body.article_id)
    if article is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Article not found")

    # Auto-assign position (next in sequence)
    existing = (await db.execute(
        select(Reference).where(Reference.paper_id == paper_id).order_by(Reference.position)
    )).scalars().all()
    position = body.position if body.position is not None else (len(existing) + 1)
    citation_text = _format_citation(article, body.citation_style, position)

    ref = Reference(
        paper_id=paper_id,
        article_id=body.article_id,
        citation_style=body.citation_style,
        citation_text=citation_text,
        position=position,
    )
    db.add(ref)
    await db.flush()
    return ref


@router.get("/{paper_id}/references", response_model=list[ReferenceResponse])
async def list_references(
    paper_id: int,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> list[Reference]:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    result = await db.execute(
        select(Reference).where(Reference.paper_id == paper_id).order_by(Reference.position)
    )
    return list(result.scalars().all())


@router.delete("/{paper_id}/references/{ref_id}", status_code=status.HTTP_204_NO_CONTENT)
async def remove_reference(
    paper_id: int,
    ref_id: int,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> None:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    ref = await db.get(Reference, ref_id)
    if ref is None or ref.paper_id != paper_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Reference not found")
    await db.delete(ref)


# ---------------------------------------------------------------------------
# AI: SECTION GENERATOR
# ---------------------------------------------------------------------------

@router.post("/{paper_id}/generate-section")
async def generate_section(
    paper_id: int,
    body: GenerateSectionRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> dict:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    import anthropic as _anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail="ANTHROPIC_API_KEY not configured")

    system_prompt = (
        "You are a medical research writing assistant specializing in peer-reviewed publications. "
        "Generate well-structured academic text suitable for a medical/health journal. "
        "Use formal language, cite provided references using numbered superscripts (e.g. [1]), "
        "and do NOT fabricate statistics, data, or citations not provided in context."
    )

    refs_text = ""
    if body.context.get("references"):
        refs_list = []
        for i, r in enumerate(body.context["references"][:20], 1):
            refs_list.append(f"[{i}] {r.get('authors', [''])[0] if r.get('authors') else ''} et al. "
                             f"{r.get('title', '')}. {r.get('journal', '')}. {r.get('year', '')}.")
        refs_text = "\n".join(refs_list)

    user_message = (
        f"Write the {body.section.upper()} section for a paper titled: \"{body.context.get('title', '')}\"\n\n"
        f"Paper abstract: {body.context.get('abstract', 'Not provided')}\n\n"
        f"Available references:\n{refs_text or 'None provided'}\n\n"
        f"Generate a complete, well-structured {body.section} section (300-600 words). "
        f"Integrate the references where appropriate."
    )

    client = _anthropic.AsyncAnthropic(api_key=api_key)
    stream = client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=1500,
        thinking={"type": "adaptive"},
        system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user_message}],
    )
    async with stream as s:
        message = await s.get_final_message()

    text = ""
    for block in message.content:
        if hasattr(block, "text"):
            text += block.text

    return {"section": body.section, "text": text.strip()}


# ---------------------------------------------------------------------------
# AI: SUGGEST REFERENCES
# ---------------------------------------------------------------------------

@router.post("/suggest-references")
async def suggest_references(
    body: SuggestReferencesRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> dict:
    """
    Full-text search within the local DB to find articles relevant to
    the provided paragraph.  (Embedding-based similarity would be added
    when a vector store is integrated.)
    """
    from sqlalchemy import func

    if len(body.paragraph_text.strip()) < 20:
        return {"suggestions": []}

    tsv = func.to_tsvector(
        "english",
        func.coalesce(Article.title, "") + " " + func.coalesce(Article.abstract, ""),
    )
    tsquery = func.plainto_tsquery("english", body.paragraph_text[:500])
    rank = func.ts_rank_cd(tsv, tsquery)

    result = await db.execute(
        select(Article)
        .where(tsv.op("@@")(tsquery))
        .order_by(rank.desc())
        .limit(body.limit)
    )
    articles = result.scalars().all()

    from app.routers.articles import ArticleResponse
    return {
        "suggestions": [
            {
                "id": a.id,
                "title": a.title,
                "authors": a.authors,
                "journal": a.journal,
                "year": a.year,
                "doi": a.doi,
                "source": a.source.value if a.source else None,
            }
            for a in articles
        ]
    }


# ---------------------------------------------------------------------------
# AI: ABSTRACT CHECKER
# ---------------------------------------------------------------------------

@router.post("/{paper_id}/check-abstract")
async def check_abstract(
    paper_id: int,
    body: CheckAbstractRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> dict:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    import anthropic as _anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail="ANTHROPIC_API_KEY not configured")

    system_prompt = (
        "You are a biomedical journal editor reviewing an abstract. "
        "Check: (1) presence of IMRAD elements (background, objective, methods, results, conclusion), "
        "(2) word count within typical 250-300 word limit, "
        "(3) consistency with provided paper content, "
        "(4) clarity and absence of jargon. "
        "Respond with JSON: {imrad_checklist: {background: bool, objective: bool, methods: bool, "
        "results: bool, conclusion: bool}, word_count: int, issues: [str], score: 0-10}"
    )

    client = _anthropic.AsyncAnthropic(api_key=api_key)
    stream = client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=800,
        system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": f"Abstract to check:\n\n{body.abstract_text}"}],
    )
    async with stream as s:
        message = await s.get_final_message()

    import re as _re
    raw = "".join(b.text for b in message.content if hasattr(b, "text"))
    # Try to parse JSON from response
    try:
        m = _re.search(r"\{.*\}", raw, _re.DOTALL)
        result = json.loads(m.group()) if m else {"raw": raw}
    except Exception:
        result = {"raw": raw}

    return result


# ---------------------------------------------------------------------------
# EXPORT
# ---------------------------------------------------------------------------

@router.get("/{paper_id}/export")
async def export_paper(
    paper_id: int,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
    format: str = Query("docx", regex="^(docx|ris|json)$"),
) -> StreamingResponse:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    if format == "json":
        content = json.dumps({
            "title": paper.title,
            "abstract": paper.abstract,
            "content": paper.content,
            "status": paper.status.value if paper.status else None,
            "doi": paper.doi,
        }, indent=2).encode()
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="paper_{paper_id}.json"'},
        )

    if format == "ris":
        # Fetch references
        refs_result = await db.execute(
            select(Reference).where(Reference.paper_id == paper_id).order_by(Reference.position)
        )
        refs = refs_result.scalars().all()

        ris_lines = ["TY  - JOUR", f"TI  - {paper.title}"]
        if paper.abstract:
            ris_lines.append(f"AB  - {paper.abstract}")
        if paper.doi:
            ris_lines.append(f"DO  - {paper.doi}")
        ris_lines.append("ER  - ")
        content = "\n".join(ris_lines).encode()
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/x-research-info-systems",
            headers={"Content-Disposition": f'attachment; filename="paper_{paper_id}.ris"'},
        )

    if format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="python-docx is not installed. Run: pip install python-docx",
            )

        doc = Document()
        doc.add_heading(paper.title or "Untitled", level=0)

        if paper.abstract:
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph(paper.abstract)

        # Extract plain text from TipTap JSON content
        def extract_text(node: dict | None) -> str:
            if node is None:
                return ""
            if node.get("type") == "text":
                return node.get("text", "")
            return "".join(extract_text(c) for c in node.get("content", []))

        if paper.content:
            for node in paper.content.get("content", []):
                node_type = node.get("type", "")
                text = extract_text(node).strip()
                if not text:
                    continue
                if node_type == "heading":
                    level = node.get("attrs", {}).get("level", 1)
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(text)

        # References section
        refs_result = await db.execute(
            select(Reference).where(Reference.paper_id == paper_id).order_by(Reference.position)
        )
        refs = refs_result.scalars().all()
        if refs:
            doc.add_heading("References", level=1)
            for ref in refs:
                doc.add_paragraph(ref.citation_text or f"[{ref.position}] Article ID {ref.article_id}")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="paper_{paper_id}.docx"'},
        )

    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported format")


# ---------------------------------------------------------------------------
# AI: LITERATURE REVIEW GENERATOR
# ---------------------------------------------------------------------------

class LiteratureReviewRequest(BaseModel):
    style: str = "narrative"   # narrative | thematic | chronological | systematic


@router.post("/{paper_id}/generate-literature-review")
async def generate_literature_review(
    paper_id: int,
    body: LiteratureReviewRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> dict:
    """
    Synthesise all references attached to this paper into a structured
    literature review section using Claude.

    Styles:
    - narrative:     Flowing prose that tells the story of the research field.
    - thematic:      Groups studies by theme/topic rather than time.
    - chronological: Ordered by publication year, tracing how the field evolved.
    - systematic:    PRISMA-aligned, with inclusion/exclusion language.
    """
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    # Fetch references with article metadata
    refs_result = await db.execute(
        select(Reference, Article)
        .join(Article, Reference.article_id == Article.id)
        .where(Reference.paper_id == paper_id)
        .order_by(Reference.position)
    )
    rows = refs_result.all()

    if not rows:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No references attached to this paper. Add references first.",
        )

    import anthropic as _anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail="ANTHROPIC_API_KEY not configured")

    # Build reference list for Claude
    ref_items: list[str] = []
    for i, (ref, article) in enumerate(rows, 1):
        authors = ", ".join((article.authors or [])[:3])
        if len(article.authors or []) > 3:
            authors += " et al."
        abstract_snippet = (article.abstract or "")[:400]
        ref_items.append(
            f"[{i}] {authors} ({article.year or 'n.d.'}). {article.title}. "
            f"{article.journal or ''}. "
            f"Study type: {article.study_type or 'unknown'}. "
            f"Abstract: {abstract_snippet}"
        )
    refs_block = "\n\n".join(ref_items)

    style_instructions = {
        "narrative": (
            "Write a flowing narrative literature review that weaves together findings "
            "from the provided studies into a coherent story of the research field. "
            "Group related findings, highlight consensus and controversy, and identify gaps."
        ),
        "thematic": (
            "Organise the literature review by themes. Identify 3-5 major themes that "
            "emerge from the provided references, then discuss each theme with supporting "
            "citations. End with a synthesis paragraph identifying research gaps."
        ),
        "chronological": (
            "Structure the literature review chronologically, tracing how understanding "
            "of the topic evolved over time. Group studies by era/decade if multiple "
            "time periods are represented. Highlight pivotal studies that changed direction."
        ),
        "systematic": (
            "Write a systematic literature review section using PRISMA-aligned language. "
            "Describe the characteristics of included studies (design, population, setting), "
            "summarise key findings, assess consistency across studies, and note limitations "
            "and quality issues. Use formal academic language."
        ),
    }.get(body.style, "Write a comprehensive literature review integrating all provided references.")

    system_prompt = (
        "You are a medical research writing expert. Generate a high-quality, academically "
        "rigorous literature review for a health/medical research paper. "
        "Cite references using numbered superscripts [1], [2], etc. "
        "Do NOT fabricate data, statistics, or citations. "
        "Only use information present in the provided references."
    )

    user_msg = (
        f"Paper title: \"{paper.title}\"\n"
        f"Paper abstract: {paper.abstract or 'Not provided'}\n\n"
        f"Literature review style: {body.style.upper()}\n"
        f"Instructions: {style_instructions}\n\n"
        f"References to synthesise ({len(rows)} papers):\n\n{refs_block}\n\n"
        f"Generate a complete literature review section (500-900 words)."
    )

    client = _anthropic.AsyncAnthropic(api_key=api_key)
    stream = client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=2000,
        system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user_msg}],
    )
    async with stream as s:
        message = await s.get_final_message()

    text = "".join(b.text for b in message.content if hasattr(b, "text"))
    return {
        "style": body.style,
        "text": text.strip(),
        "references_used": len(rows),
    }


# ---------------------------------------------------------------------------
# AI: EDIT SELECTED TEXT
# ---------------------------------------------------------------------------

class AiEditRequest(BaseModel):
    text: str
    instruction: str  # simplify | improve | paraphrase | translate | formal | custom text


@router.post("/{paper_id}/ai-edit")
async def ai_edit_text(
    paper_id: int,
    body: AiEditRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> dict:
    """
    Rewrite a selected passage according to the instruction.
    Supported instructions: simplify, improve, paraphrase, translate, formal,
    or any free-form instruction string.
    """
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    if not body.text.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No text provided")

    import anthropic as _anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail="ANTHROPIC_API_KEY not configured")

    _instruction_map = {
        "simplify":    "Simplify the following text to make it clearer and more accessible, while preserving the core meaning. Use shorter sentences and plainer language.",
        "improve":     "Improve the quality, clarity, and academic tone of the following text. Enhance word choice and sentence structure.",
        "paraphrase":  "Paraphrase the following text completely, using different wording and sentence structure while preserving the meaning.",
        "translate":   "Translate the following text to clear, formal British English.",
        "formal":      "Rewrite the following text in a more formal, academic register suitable for a peer-reviewed medical journal.",
        "shorten":     "Shorten the following text to roughly half its length while keeping the key points.",
        "expand":      "Expand the following text with more detail, evidence, and academic elaboration suitable for a medical research paper.",
    }
    instruction = _instruction_map.get(body.instruction.lower(), body.instruction)

    system_prompt = (
        "You are an expert medical research writing assistant. "
        "Rewrite the provided text exactly as instructed. "
        "Return ONLY the rewritten text with no preamble, explanation, or quotation marks."
    )

    user_msg = f"{instruction}\n\nText to rewrite:\n\n{body.text}"

    client = _anthropic.AsyncAnthropic(api_key=api_key)
    stream = client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=2000,
        system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user_msg}],
    )
    async with stream as s:
        message = await s.get_final_message()

    text = "".join(b.text for b in message.content if hasattr(b, "text"))
    return {"text": text.strip(), "instruction": body.instruction}


# ---------------------------------------------------------------------------
# AI: INLINE AUTOCOMPLETE
# ---------------------------------------------------------------------------

class AutocompleteRequest(BaseModel):
    text: str     # The text the user has written so far in the section
    section: str  # Section heading for context (e.g. "Introduction")
    max_tokens: int = 150


@router.post("/{paper_id}/autocomplete")
async def autocomplete(
    paper_id: int,
    body: AutocompleteRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> dict:
    """
    Continue the user's writing from where they left off.
    Returns a short completion (1-3 sentences) to be shown as a suggestion.
    """
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    # Only trigger on non-trivial text
    if len(body.text.strip()) < 20:
        return {"completion": ""}

    import anthropic as _anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        return {"completion": ""}

    system_prompt = (
        "You are an AI writing assistant integrated into a medical research paper editor. "
        "Continue the researcher's writing naturally and academically. "
        "Write exactly 1-3 sentences that flow seamlessly from the provided text. "
        "Do NOT repeat any text already written. Do NOT add a heading. "
        "Return only the continuation text."
    )

    # Trim to last 1500 chars to stay within context
    context = body.text[-1500:].strip()
    user_msg = (
        f"Paper section: {body.section}\n\n"
        f"Text so far:\n{context}\n\n"
        "Continue writing:"
    )

    client = _anthropic.AsyncAnthropic(api_key=api_key)
    stream = client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=min(body.max_tokens, 300),
        system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user_msg}],
    )
    async with stream as s:
        message = await s.get_final_message()

    completion = "".join(b.text for b in message.content if hasattr(b, "text"))
    return {"completion": completion.strip()}


# ---------------------------------------------------------------------------
# DOI MINTING via Zenodo
# ---------------------------------------------------------------------------

class PublishResponse(BaseModel):
    doi: str
    paper_id: int


@router.post(
    "/{paper_id}/publish",
    response_model=PublishResponse,
    summary="Mint a DOI via Zenodo and mark paper as published",
)
async def publish_paper(
    paper_id: int,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: CurrentUser,
) -> PublishResponse:
    paper = await db.get(Paper, paper_id)
    if paper is None or paper.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    if paper.doi:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"Paper already has DOI: {paper.doi}",
        )

    from app.services.zenodo import mint_doi

    try:
        doi = await mint_doi(
            title=paper.title,
            abstract=paper.abstract,
            content=paper.content,
            author_name=current_user.name,
        )
    except RuntimeError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        )

    paper.doi = doi
    paper.status = PaperStatus.PUBLISHED

    return PublishResponse(doi=doi, paper_id=paper_id)
